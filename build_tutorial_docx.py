# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── 页面设置 A4 ──────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = section.right_margin = Cm(2.5)
section.top_margin  = section.bottom_margin = Cm(2.5)

# ── 字体默认中文宋体 ──────────────────────────────────────
style_normal = doc.styles['Normal']
style_normal.font.name = '宋体'
style_normal.font.size = Pt(11)
style_normal._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ── 工具函数 ─────────────────────────────────────────────
def add_heading(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        else:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    return p

def add_para(text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(11)
    run.bold = bold
    return p

def add_code(lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  'F2F2F2')
        pPr.append(shd)
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x6E)
    doc.add_paragraph()

def add_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.name = '黑体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                run.font.size = Pt(10)
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  'BDD7EE')
        tcPr.append(shd)
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx+1].cells
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = str(val)
            for para in row_cells[c_idx].paragraphs:
                for run in para.runs:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(10)
    doc.add_paragraph()

def add_note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'FFF2CC')
    pPr.append(shd)
    run = p.add_run('提示：' + text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x7F, 0x60, 0x00)
    doc.add_paragraph()

# ════════════════════════════════════════════════════════════
# 封面
# ════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('\n\n\n\n')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('催化剂产率预测\nXGBoost 机器学习完整教程')
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('基于 NeSI 高性能计算平台\n面向机器学习初学者')
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 文档结构总览
# ════════════════════════════════════════════════════════════
add_heading('文档结构总览', 1)
add_para('本教程分为以下章节，建议按顺序操作：')
add_table(['章节', '标题', '内容'], [
    ('第0章', '项目结构与准备工作', '在本地创建5个文件夹，整体上传到NeSI'),
    ('第1章', '环境配置',           '用 venv 创建隔离的 Python 环境并安装依赖'),
    ('第2章', '读取数据',           '从 data/ 加载 Excel，理解数据结构'),
    ('第3章', '数据预处理',         '分离特征与目标，编码分类变量'),
    ('第4章', '训练模型',           'XGBoost 训练、评估，模型保存到 models/'),
    ('第5章', '特征重要性',         '分析配方因素影响，图片保存到 logs/'),
    ('第6章', 'SLURM 作业提交',    '写作业脚本，日志输出到 logs/，正式提交'),
    ('附录',  '完整源代码',         'src/train_xgb_lgb.py 完整内容'),
])

# ════════════════════════════════════════════════════════════
# 第0章：项目结构
# ════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading('第0章：项目结构与准备工作', 1)

add_heading('0.1 项目文件夹结构', 2)
add_para('你在本地（Windows）维护如下5个文件夹，整体上传到 NeSI，结构保持一致：')
add_code([
    'FCC-AI/                         <- 项目根目录',
    '├── data/                       <- 存放输入数据文件',
    '│   └── ml_A_all.xlsx  <- 本项目使用的数据',
    '├── logs/                       <- 存放运行日志和图片输出',
    '│   ├── xgb_log_<任务ID>.txt   <- SLURM 标准输出日志（自动生成）',
    '│   └── importance_*.png        <- 特征重要性图（训练后生成）',
    '├── models/                     <- 存放训练好的模型文件',
    '│   └── model_C4_汽油.json      <- 各产率的模型（训练后生成）',
    '├── src/                        <- 存放所有 Python 源代码',
    '│   ├── train_xgb_lgb.py           <- 主训练脚本（本教程核心）',
    '│   └── run_xgb.sl             <- SLURM 作业提交脚本',
    '└── venv/                       <- Python 虚拟环境（在NeSI上创建）',
])
add_note('venv/ 文件夹在 NeSI 上创建即可，不需要从本地上传（体积很大）。其余4个文件夹从本地整体上传。')

add_heading('0.2 在本地准备文件夹', 2)
add_para('在 Windows 上打开 PowerShell，创建项目结构：')
add_code([
    '# 在桌面或任意位置创建项目根目录',
    'mkdir FCC-AI',
    'cd FCC-AI',
    'mkdir data logs models src',
    '',
    '# 把数据文件放到 data/ 里',
    'copy "D:\\cc_test\\ml_A_all.xlsx" data\\',
    '',
    '# 把写好的源代码放到 src/ 里（见附录）',
    '# 把 train_xgb_lgb.py 和 run_xgb.sl 复制到 src\\',
])

add_heading('0.3 登录 NeSI', 2)
add_para('在 Windows PowerShell 中执行（将 your_username 替换为你的账号）：')
add_code([
    '# 第一步：登录跳板节点',
    'ssh your_username@lander.nesi.org.nz',
    '',
    '# 第二步：进入计算集群',
    'ssh mahuika',
    '',
    '# 登录成功后，确认当前目录',
    'pwd',
    '# 输出类似：/home/your_username',
])

add_heading('0.4 将整个项目文件夹上传到 NeSI', 2)
add_para('在本地另开一个 PowerShell 窗口（不要关闭 NeSI 连接），运行：')
add_code([
    '# 用 scp -r 上传整个文件夹（-r 表示递归上传子目录）',
    '# 说明：在本地 FCC-AI 目录内执行以下命令',
    '# venv/ 不需要上传，在 NeSI 上重新创建即可',
    '# 在本地 FCC-AI 文件夹内执行（cd 到 FCC-AI 目录后运行）',
    'scp -r data  your_username@lander.nesi.org.nz:/nesi/project/uoa04367/PROJECT-04367/FCC-AI/',
    'scp -r logs  your_username@lander.nesi.org.nz:/nesi/project/uoa04367/PROJECT-04367/FCC-AI/',
    'scp -r models your_username@lander.nesi.org.nz:/nesi/project/uoa04367/PROJECT-04367/FCC-AI/',
    'scp -r src   your_username@lander.nesi.org.nz:/nesi/project/uoa04367/PROJECT-04367/FCC-AI/',
    '',
    '# 或者一次性上传（需要先在 NeSI 上创建目录（已存在则跳过））',
    '# 在 NeSI 上：',
    'mkdir -p /nesi/project/uoa04367/PROJECT-04367/FCC-AI/{data,logs,models,src,venv}',
    '',
    '# 在本地（上传4个文件夹内容）：',
    'scp data/ml_A_all.xlsx your_username@lander.nesi.org.nz:/nesi/project/uoa04367/PROJECT-04367/FCC-AI/data/',
    'scp src/train_xgb_lgb.py           your_username@lander.nesi.org.nz:/nesi/project/uoa04367/PROJECT-04367/FCC-AI/src/',
    'scp src/run_xgb.sl             your_username@lander.nesi.org.nz:/nesi/project/uoa04367/PROJECT-04367/FCC-AI/src/',
])
add_note('logs/ 和 models/ 可以是空文件夹，脚本运行时会自动往里写入结果。')

add_heading('0.5 在 NeSI 上确认文件结构', 2)
add_code([
    '# 在 NeSI 上查看目录结构',
    'ls -R /nesi/project/uoa04367/PROJECT-04367/FCC-AI/',
    '',
    '# 确认数据文件在位',
    'ls /nesi/project/uoa04367/PROJECT-04367/FCC-AI/data/',
    '# 应看到：ml_A_all.xlsx',
    '',
    '# 确认源码在位',
    'ls /nesi/project/uoa04367/PROJECT-04367/FCC-AI/src/',
    '# 应看到：train_xgb_lgb.py  run_xgb.sl',
])

# ════════════════════════════════════════════════════════════
# 第1章：环境配置（venv）
# ════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading('第1章：环境配置 —— 在 NeSI 上创建 venv 虚拟环境', 1)

add_heading('1.1 什么是 venv 虚拟环境', 2)
add_para(
    'venv 是 Python 内置的虚拟环境工具。它会在你的 venv/ 文件夹中创建一个'
    '独立的 Python 环境，所有安装的包都只属于这个项目，不会影响 NeSI 的系统环境，'
    '也不会和其他项目的包版本冲突。这与你本地的 venv/ 文件夹的用途完全一致。'
)

add_heading('1.2 加载 Python 模块并创建 venv', 2)
add_para('在 NeSI 上执行（只需做一次）：')
add_code([
    '# 进入项目根目录',
    'cd /nesi/project/uoa04367/PROJECT-04367/FCC-AI',
    '',
    '# 加载 NeSI 提供的 Python',
    'module purge',
    'module load Python/3.10.5-gimkl-2022a',
    '',
    '# 在 venv/ 文件夹中创建虚拟环境',
    'python -m venv venv',
    '',
    '# 激活虚拟环境（每次登录后都要执行这一步）',
    'source venv/bin/activate',
    '',
    '# 激活成功后，命令行前面会出现 (venv) 提示符',
    '# 例如：(venv) [your_username@mahuika FCC-AI]$',
])

add_heading('1.3 安装所需 Python 包', 2)
add_para('激活 venv 后，安装机器学习所需的全部包（只需安装一次）：')
add_code([
    '# 确保 venv 已激活（看到 (venv) 前缀）',
    'pip install xgboost pandas openpyxl matplotlib scikit-learn',
    '',
    '# 安装完成后验证',
    'python -c "import xgboost; print(\'xgboost 版本:\', xgboost.__version__)"',
    'python -c "import pandas;   print(\'pandas 版本:\',  pandas.__version__)"',
])
add_para('各包作用说明：')
add_table(['包名', '作用'], [
    ('xgboost',      'XGBoost 模型，本教程核心'),
    ('pandas',       '读取 Excel、处理表格数据'),
    ('openpyxl',     'pandas 读取 .xlsx 的底层依赖'),
    ('matplotlib',   '绘制特征重要性图，保存到 logs/'),
    ('scikit-learn', '数据分割、评估指标（R²/MAE/RMSE）'),
])

add_heading('1.4 之后每次登录的操作', 2)
add_para('每次重新登录 NeSI 后，需要重新激活 venv（模块也需重新加载）：')
add_code([
    'cd /nesi/project/uoa04367/PROJECT-04367/FCC-AI',
    'module purge',
    'module load Python/3.10.5-gimkl-2022a',
    'source venv/bin/activate',
    '# 看到 (venv) 前缀后即可开始工作',
])
add_note('SLURM 作业脚本中会自动执行上面的激活步骤，所以提交任务时不需要手动激活。')

# ════════════════════════════════════════════════════════════
# 第2章：读取数据
# ════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading('第2章：读取数据 —— 从 data/ 加载 Excel', 1)

add_heading('2.1 脚本位置和路径约定', 2)
add_para('所有 Python 脚本放在 src/ 目录，使用相对路径读取 data/ 中的数据：')
add_code([
    'FCC-AI/   （即 /nesi/project/uoa04367/PROJECT-04367/FCC-AI）',
    '├── data/ml_A_all.xlsx   <- 数据在这里',
    '└── src/train_xgb_lgb.py             <- 脚本从这里读数据，路径写 ../data/...',
])

add_heading('2.2 数据读取代码（train_xgb_lgb.py 开头部分）', 2)
add_code([
    '# -*- coding: utf-8 -*-',
    'import pandas as pd',
    'import numpy as np',
    'import matplotlib',
    'matplotlib.use("Agg")           # NeSI 无显示器，图片保存为文件',
    'import matplotlib.pyplot as plt',
    'import os',
    'from pathlib import Path',
    '',
    '# ── 路径设置 ──────────────────────────────────────────',
    '# __file__ 是当前脚本自身的路径（src/train_xgb_lgb.py）',
    '# .parent 是 src/ 目录，再 .parent 是项目根目录',
    'ROOT = Path(__file__).resolve().parent.parent',
    '',
    'DATA_DIR   = ROOT / "data"',
    'LOGS_DIR   = ROOT / "logs"',
    'MODELS_DIR = ROOT / "models"',
    '',
    '# 确保输出目录存在',
    'LOGS_DIR.mkdir(exist_ok=True)',
    'MODELS_DIR.mkdir(exist_ok=True)',
    '',
    '# ── 读取数据 ──────────────────────────────────────────',
    'data_file = DATA_DIR / "ml_A_all.xlsx"',
    'print(f"读取数据: {data_file}")',
    'df = pd.read_excel(data_file, sheet_name="ML_data")',
    'print(f"数据维度: {df.shape[0]} 行 x {df.shape[1]} 列")',
    '',
    '# ── 查看基本信息 ──────────────────────────────────────',
    'print("\\n各列缺失值统计:")',
    'print(df.isnull().sum().to_string())',
])

add_heading('2.3 路径代码解释', 2)
add_table(['代码', '解释'], [
    ('Path(__file__).resolve()',  '获取当前脚本 train_xgb_lgb.py 的绝对路径'),
    ('.parent.parent',            '从 src/ 上两级到项目根目录 FCC-AI/'),
    ('ROOT / "data"',             '拼接路径，等价于 FCC-AI/data/'),
    ('mkdir(exist_ok=True)',      '如果目录不存在则创建，已存在则不报错'),
    ('pd.read_excel(...)',         '读取 Excel 文件，sheet_name 指定读 ML_data 表'),
])

# ════════════════════════════════════════════════════════════
# 第3章：数据预处理
# ════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading('第3章：数据预处理 —— 分离特征与目标，编码原料油', 1)

add_heading('3.1 特征列与目标列的区分', 2)
add_para('数据中所有以 C 开头的列是目标变量（产率），其余列是特征（输入）：')
add_para('• 特征列（X）：配方（M/FM/Z/FZ）+ 理化性质（L）+ 操作条件（T1/T2）+ 原料油（oil）', indent=True)
add_para('• 目标列（Y）：C1_焦炭、C2_干气、C3_液化气、C4_汽油、C9_转化率 等共24列', indent=True)

add_heading('3.2 预处理代码', 2)
add_code([
    'from sklearn.preprocessing import LabelEncoder',
    '',
    '# ── 分离特征和目标 ────────────────────────────────',
    'target_cols  = [c for c in df.columns if c.startswith("C")]',
    'feature_cols = [c for c in df.columns if not c.startswith("C")]',
    '',
    'X = df[feature_cols].copy()    # 特征矩阵（31列）',
    'Y = df[target_cols].copy()     # 目标矩阵（24列）',
    '',
    'print(f"特征列: {len(feature_cols)} 个")',
    'print(f"目标列: {len(target_cols)} 个")',
    '',
    '# ── 编码原料油（文字 → 数字） ─────────────────────',
    '# oil_原料油 列是字符串（如 "300W"），XGBoost 只接受数字',
    'le = LabelEncoder()',
    'X["oil_原料油"] = X["oil_原料油"].fillna("未知")',
    'X["oil_原料油"] = le.fit_transform(X["oil_原料油"])',
    '',
    'print("\\n原料油编码对照：")',
    'for i, name in enumerate(le.classes_):',
    '    print(f"  {name} -> {i}")',
    '',
    '# ── 其他 NaN 无需处理 ─────────────────────────────',
    '# XGBoost 原生支持 NaN，配方列和理化性质列的空值',
    '# 模型会自动学习最优的分裂方向，无需手动填充',
    'print("\\n预处理完成，特征矩阵维度:", X.shape)',
])

# ════════════════════════════════════════════════════════════
# 第4章：训练模型
# ════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading('第4章：训练 XGBoost 模型并保存到 models/', 1)

add_heading('4.1 什么是 XGBoost', 2)
add_para(
    'XGBoost（Extreme Gradient Boosting）是一种集成学习算法，通过将大量简单的'
    '决策树逐一叠加来做出准确预测。每棵新树专门纠正前面所有树的错误，最终组合出'
    '强大的预测能力。它的优势：原生支持缺失值、对稀疏配方数据友好、可输出特征重要性。'
)

add_heading('4.2 训练代码（以 C4_汽油 为例）', 2)
add_code([
    'from sklearn.model_selection import train_test_split',
    'from sklearn.metrics import r2_score, mean_absolute_error',
    'from xgboost import XGBRegressor',
    '',
    '# ── 选择目标变量 ─────────────────────────────────',
    'target_name = "C4_汽油"      # 修改这里可切换预测目标',
    'y = Y[target_name]',
    '',
    '# ── 过滤掉目标值为空的行 ──────────────────────────',
    'valid = y.notna()',
    'X_use = X[valid]',
    'y_use = y[valid]',
    'print(f"有效样本数: {len(y_use)}")',
    '',
    '# ── 划分训练集（80%）和测试集（20%） ────────────',
    'X_train, X_test, y_train, y_test = train_test_split(',
    '    X_use, y_use, test_size=0.2, random_state=42',
    ')',
    'print(f"训练集: {len(X_train)} 行  |  测试集: {len(X_test)} 行")',
    '',
    '# ── 定义模型 ─────────────────────────────────────',
    'model = XGBRegressor(',
    '    n_estimators=300,      # 决策树数量',
    '    learning_rate=0.05,    # 学习率',
    '    max_depth=6,           # 每棵树的最大深度',
    '    subsample=0.8,         # 每棵树随机使用80%样本（防止过拟合）',
    '    colsample_bytree=0.8,  # 每棵树随机使用80%特征',
    '    random_state=42,',
    '    n_jobs=-1,             # 用全部CPU核心',
    '    verbosity=0',
    ')',
    '',
    '# ── 训练模型 ─────────────────────────────────────',
    'model.fit(X_train, y_train,',
    '          eval_set=[(X_test, y_test)],',
    '          verbose=False)',
    'print("训练完成！")',
    '',
    '# ── 评估 ─────────────────────────────────────────',
    'y_pred = model.predict(X_test)',
    'r2   = r2_score(y_test, y_pred)',
    'mae  = mean_absolute_error(y_test, y_pred)',
    'rmse = float(np.sqrt(np.mean((y_test - y_pred)**2)))',
    'print(f"R2  = {r2:.4f}  (越接近1越好，>0.8为良好)")',
    'print(f"MAE = {mae:.4f}  (平均绝对误差，单位为%)")',
    'print(f"RMSE= {rmse:.4f}  (均方根误差，单位为%)")',
    '',
    '# ── 保存模型到 models/ ────────────────────────────',
    'model_path = MODELS_DIR / f"model_{target_name}.json"',
    'model.save_model(model_path)',
    'print(f"模型已保存: {model_path}")',
    '',
    '# ── 保存预测结果到 logs/ ──────────────────────────',
    'result_df = pd.DataFrame({',
    '    "实际值": y_test.values,',
    '    "预测值": y_pred,',
    '    "误差":   y_test.values - y_pred',
    '})',
    'result_path = LOGS_DIR / f"result_{target_name}.csv"',
    'result_df.to_csv(result_path, index=False)',
    'print(f"预测结果已保存: {result_path}")',
])

add_heading('4.3 模型参数说明', 2)
add_table(['参数', '含义', '建议值'], [
    ('n_estimators',     '决策树数量，越多越精细但训练越慢',     '200~500，从300开始'),
    ('learning_rate',    '每棵树的学习步长，越小需要越多的树',    '0.05~0.1'),
    ('max_depth',        '每棵树最大层数，越深模型越复杂',       '4~8，过深容易过拟合'),
    ('subsample',        '每棵树随机取样的比例',                 '0.7~1.0'),
    ('colsample_bytree', '每棵树随机选取的特征比例',             '0.7~1.0'),
    ('n_jobs=-1',        '使用全部可用CPU核心并行计算',          '固定-1即可'),
])

add_heading('4.4 评估指标说明', 2)
add_table(['指标', '含义', '判断标准'], [
    ('R²',   '决定系数：模型解释了多少数据变化，1=完美，0=等同于猜均值', '>0.85优秀，>0.7良好，<0.5需改进'),
    ('MAE',  '平均绝对误差：预测值与真实值的平均偏差，单位为%',          '越小越好，<1% 为优秀'),
    ('RMSE', '均方根误差：对大偏差更敏感，同样单位为%',                  '越小越好，通常略大于MAE'),
])

# ════════════════════════════════════════════════════════════
# 第5章：特征重要性
# ════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading('第5章：特征重要性分析 —— 图片保存到 logs/', 1)

add_heading('5.1 特征重要性的意义', 2)
add_para(
    '训练完成后，XGBoost 可以告诉你每个特征（配方组分、反应温度等）'
    '对预测结果的贡献度。重要性高的特征对产率影响大，'
    '这对于指导配方优化具有直接的实用价值。'
)

add_heading('5.2 特征重要性代码', 2)
add_code([
    '# ── 整理特征重要性 ───────────────────────────────',
    'imp_df = pd.DataFrame({',
    '    "特征":  X_use.columns.tolist(),',
    '    "重要性": model.feature_importances_',
    '}).sort_values("重要性", ascending=False)',
    '',
    '# 打印 Top10',
    'print("\\n===== 特征重要性 Top10 =====")',
    'print(imp_df.head(10).to_string(index=False))',
    '',
    '# ── 保存重要性数据到 logs/ ───────────────────────',
    'imp_csv = LOGS_DIR / f"importance_{target_name}.csv"',
    'imp_df.to_csv(imp_csv, index=False)',
    '',
    '# ── 绘图并保存到 logs/ ───────────────────────────',
    'fig, ax = plt.subplots(figsize=(10, 7))',
    'top = imp_df.head(15)',
    'ax.barh(top["特征"][::-1], top["重要性"][::-1], color="#2E74B5")',
    'ax.set_xlabel("重要性得分", fontsize=12)',
    'ax.set_title(f"{target_name} - Top15 特征重要性", fontsize=14)',
    'plt.tight_layout()',
    'img_path = LOGS_DIR / f"importance_{target_name}.png"',
    'plt.savefig(img_path, dpi=150)',
    'plt.close()',
    'print(f"特征重要性图已保存: {img_path}")',
])
add_note('运行结束后，从 NeSI 下载 logs/*.png 到本地查看图片效果。')

# ════════════════════════════════════════════════════════════
# 第6章：SLURM 作业提交
# ════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading('第6章：SLURM 作业提交 —— 在 NeSI 上运行任务', 1)

add_heading('6.1 什么是 SLURM', 2)
add_para(
    'NeSI 是多用户共享平台，不能直接运行耗时程序。'
    'SLURM 是任务调度系统：你提交一个"作业脚本"说明需要多少资源，'
    '系统分配到空闲节点后自动运行你的代码，日志写入 logs/ 目录。'
)

add_heading('6.2 SLURM 作业脚本（src/run_xgb.sl）', 2)
add_para('在 src/ 目录下创建 run_xgb.sl 文件，内容如下：')
add_code([
    '#!/bin/bash',
    '#SBATCH --job-name=xgb_catalyst          # 任务名称',
    '#SBATCH --account=your_project_code      # 你的NeSI项目代码（必填）',
    '#SBATCH --time=02:00:00                  # 最长运行2小时',
    '#SBATCH --mem=16GB                       # 申请16GB内存',
    '#SBATCH --cpus-per-task=8               # 申请8个CPU核心',
    '#SBATCH --output=../logs/xgb_log_%j.txt # 标准输出 -> logs/目录',
    '#SBATCH --error=../logs/xgb_err_%j.txt  # 错误日志 -> logs/目录',
    '',
    '# ── 进入项目根目录 ────────────────────────────────',
    'cd /nesi/project/uoa04367/PROJECT-04367/FCC-AI',
    '',
    '# ── 加载 Python 模块 ──────────────────────────────',
    'module purge',
    'module load Python/3.10.5-gimkl-2022a',
    '',
    '# ── 激活 venv 虚拟环境 ────────────────────────────',
    'source venv/bin/activate',
    '',
    '# ── 运行主脚本 ────────────────────────────────────',
    'python src/train_xgb_lgb.py',
    '',
    'echo "==============================="',
    'echo "任务完成！"',
    'echo "日志和图片在 logs/ 目录"',
    'echo "模型文件在 models/ 目录"',
])

add_heading('6.3 SLURM 参数说明', 2)
add_table(['参数', '含义', '注意事项'], [
    ('--account',       'NeSI 项目代码',              '必须填写，登录后用 id 命令查看，或问管理员'),
    ('--time',          '任务最长运行时间',             '超时会被强制停止，宁可设长一些'),
    ('--mem',           '申请的内存',                  '我们数据量小，16GB 绰绰有余'),
    ('--cpus-per-task', '申请的 CPU 核心数',           'XGBoost 的 n_jobs=-1 会用满所有核'),
    ('--output',        'SLURM 标准输出写入的文件路径', '用 ../logs/ 让日志自动存入 logs/ 目录'),
    ('--error',         '错误信息写入的文件路径',       '出错时在这里找原因'),
])

add_heading('6.4 提交任务', 2)
add_code([
    '# 在 NeSI 上，进入 src/ 目录提交',
    'cd /nesi/project/uoa04367/PROJECT-04367/FCC-AI/src',
    'sbatch run_xgb.sl',
    '# 成功提交后输出：Submitted batch job 12345678',
])

add_heading('6.5 管理和监控任务', 2)
add_code([
    '# 查看自己的任务状态',
    '# ST 列：R=运行中  PD=排队等待  CG=即将结束',
    'squeue --me',
    '',
    '# 实时追踪日志输出（Ctrl+C 停止追踪，任务继续运行）',
    'tail -f /nesi/project/uoa04367/PROJECT-04367/FCC-AI/logs/xgb_log_12345678.txt',
    '',
    '# 取消任务（如果发现参数设置错了）',
    'scancel 12345678',
    '',
    '# 任务完成后，查看输出文件',
    'ls /nesi/project/uoa04367/PROJECT-04367/FCC-AI/logs/',
    'ls /nesi/project/uoa04367/PROJECT-04367/FCC-AI/models/',
])

add_heading('6.6 将结果下载到本地', 2)
add_para('任务完成后，在本地 PowerShell 中运行：')
add_code([
    '# 创建本地接收目录',
    'mkdir D:\\cc_test\\results',
    '',
    '# 下载 logs/ 中的日志、图片和 CSV',
    'scp -r your_username@lander.nesi.org.nz:/nesi/project/uoa04367/PROJECT-04367/FCC-AI/logs/ D:\\cc_test\\results\\',
    '',
    '# 下载 models/ 中的模型文件',
    'scp -r your_username@lander.nesi.org.nz:/nesi/project/uoa04367/PROJECT-04367/FCC-AI/models/ D:\\cc_test\\results\\',
])

# ════════════════════════════════════════════════════════════
# 附录：完整 Python 脚本
# ════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading('附录：完整源代码 —— src/train_xgb_lgb.py', 1)
add_para('将以下全部代码保存为 src/train_xgb_lgb.py，上传到 NeSI 后即可通过 run_xgb.sl 提交运行。')
add_para('脚本对每个产率目标同时训练 XGBoost 和 LightGBM，自动比较精度，结果保存到 logs/，模型保存到 models/。')

add_code([
    '# -*- coding: utf-8 -*-',
    '"""',
    'FCC 催化剂产率预测 - XGBoost + LightGBM 双模型对比',
    '项目结构：',
    '  /nesi/project/uoa04367/PROJECT-04367/FCC-AI/',
    '  ├── data/ml_A_all.xlsx  <- 输入数据',
    '  ├── logs/                        <- 日志、图片、CSV 输出',
    '  ├── models/                      <- 模型文件输出',
    '  ├── src/train_xgb_lgb.py         <- 本脚本',
    '  └── venv/                        <- Python 虚拟环境',
    '"""',
    '',
    '# ── 导入库 ──────────────────────────────────────────────',
    'import pandas as pd',
    'import numpy as np',
    'import matplotlib',
    'matplotlib.use("Agg")',
    'import matplotlib.pyplot as plt',
    'from pathlib import Path',
    'from sklearn.model_selection import train_test_split',
    'from sklearn.metrics import r2_score, mean_absolute_error',
    'from sklearn.preprocessing import LabelEncoder',
    'from xgboost import XGBRegressor',
    'from lightgbm import LGBMRegressor',
    '',
    '# ── 路径配置 ─────────────────────────────────────────────',
    'ROOT       = Path(__file__).resolve().parent.parent',
    'DATA_DIR   = ROOT / "data"',
    'LOGS_DIR   = ROOT / "logs"',
    'MODELS_DIR = ROOT / "models"',
    'LOGS_DIR.mkdir(exist_ok=True)',
    'MODELS_DIR.mkdir(exist_ok=True)',
    '',
    '# ── 1. 读取数据 ──────────────────────────────────────────',
    'print("=" * 60)',
    'print("FCC 催化剂产率预测 - XGBoost + LightGBM 双模型")',
    'print("=" * 60)',
    'df = pd.read_excel(DATA_DIR / "ml_A_all.xlsx", sheet_name="ML_data")',
    'print(f"数据维度: {df.shape[0]} 行 x {df.shape[1]} 列")',
    '',
    '# ── 2. 分离特征和目标 ────────────────────────────────────',
    'target_cols  = [c for c in df.columns if c.startswith("C")]',
    'feature_cols = [c for c in df.columns if not c.startswith("C")]',
    'X = df[feature_cols].copy()',
    'Y = df[target_cols].copy()',
    'print(f"特征: {len(feature_cols)} 列  |  目标: {len(target_cols)} 列")',
    '',
    '# ── 3. 编码原料油 ────────────────────────────────────────',
    'le = LabelEncoder()',
    'X["oil_原料油"] = X["oil_原料油"].fillna("未知")',
    'X["oil_原料油"] = le.fit_transform(X["oil_原料油"])',
    '',
    '# ── 4. 循环训练每个目标 ──────────────────────────────────',
    'summary = []',
    '',
    '# 优先训练这些核心目标',
    'priority = ["C4_汽油", "C9_转化率", "C3_液化气",',
    '            "C1_焦炭", "C5_柴油",   "C8_总液收"]',
    'all_targets = priority + [c for c in target_cols if c not in priority]',
    '',
    'for target_name in all_targets:',
    '    print(f"\\n--- 训练目标: {target_name} ---")',
    '',
    '    y = Y[target_name]',
    '    valid = y.notna()',
    '    X_use, y_use = X[valid], y[valid]',
    '',
    '    if len(y_use) < 50:',
    '        print(f"  样本不足 ({len(y_use)})，跳过")',
    '        continue',
    '',
    '    X_train, X_test, y_train, y_test = train_test_split(',
    '        X_use, y_use, test_size=0.2, random_state=42)',
    '',
    '    # ── XGBoost ──────────────────────────────────────────',
    '    xgb = XGBRegressor(',
    '        n_estimators=300, learning_rate=0.05, max_depth=6,',
    '        subsample=0.8, colsample_bytree=0.8,',
    '        random_state=42, n_jobs=-1, verbosity=0)',
    '    xgb.fit(X_train, y_train,',
    '            eval_set=[(X_test, y_test)], verbose=False)',
    '    y_xgb    = xgb.predict(X_test)',
    '    r2_xgb   = r2_score(y_test, y_xgb)',
    '    mae_xgb  = mean_absolute_error(y_test, y_xgb)',
    '    rmse_xgb = float(np.sqrt(np.mean((y_test - y_xgb) ** 2)))',
    '    print(f"  XGB  R2={r2_xgb:.4f}  MAE={mae_xgb:.4f}  RMSE={rmse_xgb:.4f}")',
    '',
    '    # ── LightGBM ─────────────────────────────────────────',
    '    lgb = LGBMRegressor(',
    '        n_estimators=300, learning_rate=0.05, max_depth=6,',
    '        subsample=0.8, colsample_bytree=0.8,',
    '        random_state=42, n_jobs=-1, verbose=-1)',
    '    lgb.fit(X_train, y_train)',
    '    y_lgb    = lgb.predict(X_test)',
    '    r2_lgb   = r2_score(y_test, y_lgb)',
    '    mae_lgb  = mean_absolute_error(y_test, y_lgb)',
    '    rmse_lgb = float(np.sqrt(np.mean((y_test - y_lgb) ** 2)))',
    '    print(f"  LGB  R2={r2_lgb:.4f}  MAE={mae_lgb:.4f}  RMSE={rmse_lgb:.4f}")',
    '',
    '    best = "XGB" if r2_xgb >= r2_lgb else "LGB"',
    '    print(f"  最佳模型: {best}")',
    '',
    '    # 保存模型',
    '    xgb.save_model(MODELS_DIR / f"model_xgb_{target_name}.json")',
    '    lgb.booster_.save_model(str(MODELS_DIR / f"model_lgb_{target_name}.txt"))',
    '',
    '    # 保存预测结果',
    '    pd.DataFrame({',
    '        "实际值":    y_test.values,',
    '        "XGB预测值": y_xgb,',
    '        "LGB预测值": y_lgb,',
    '    }).to_csv(LOGS_DIR / f"result_{target_name}.csv", index=False)',
    '',
    '    # 特征重要性（XGB）',
    '    imp_xgb = pd.DataFrame({',
    '        "特征":      X_use.columns.tolist(),',
    '        "XGB重要性": xgb.feature_importances_,',
    '    }).sort_values("XGB重要性", ascending=False)',
    '    imp_xgb.to_csv(LOGS_DIR / f"importance_xgb_{target_name}.csv", index=False)',
    '',
    '    # 特征重要性（LGB）',
    '    imp_lgb = pd.DataFrame({',
    '        "特征":      X_use.columns.tolist(),',
    '        "LGB重要性": lgb.feature_importances_,',
    '    }).sort_values("LGB重要性", ascending=False)',
    '    imp_lgb.to_csv(LOGS_DIR / f"importance_lgb_{target_name}.csv", index=False)',
    '',
    '    # 重点目标：XGB vs LGB 双图对比',
    '    if target_name in priority:',
    '        fig, axes = plt.subplots(1, 2, figsize=(18, 7))',
    '        for ax, imp, label, color in zip(',
    '            axes,',
    '            [imp_xgb.rename(columns={"XGB重要性": "重要性"}),',
    '             imp_lgb.rename(columns={"LGB重要性": "重要性"})],',
    '            ["XGBoost", "LightGBM"],',
    '            ["#2E74B5", "#E87722"],',
    '        ):',
    '            top = imp.head(15)',
    '            ax.barh(top["特征"][::-1], top["重要性"][::-1], color=color)',
    '            ax.set_xlabel("重要性得分")',
    '            ax.set_title(f"{target_name} - {label} Top15 特征重要性")',
    '        plt.tight_layout()',
    '        plt.savefig(LOGS_DIR / f"importance_{target_name}.png", dpi=150)',
    '        plt.close()',
    '',
    '    summary.append({',
    '        "目标变量":   target_name,',
    '        "样本数":     len(y_use),',
    '        "XGB_R2":     round(r2_xgb,  4),',
    '        "XGB_MAE":    round(mae_xgb,  4),',
    '        "XGB_RMSE":   round(rmse_xgb, 4),',
    '        "LGB_R2":     round(r2_lgb,  4),',
    '        "LGB_MAE":    round(mae_lgb,  4),',
    '        "LGB_RMSE":   round(rmse_lgb, 4),',
    '        "最佳模型":   best,',
    '    })',
    '',
    '# ── 5. 汇总结果 ──────────────────────────────────────────',
    'summary_df = pd.DataFrame(summary)',
    'summary_df.to_csv(LOGS_DIR / "summary_xgb_vs_lgb.csv", index=False)',
    '',
    'print("\\n" + "=" * 60)',
    'print("全部完成！")',
    'print(f"  模型文件 -> {MODELS_DIR}")',
    'print(f"  图片/CSV -> {LOGS_DIR}")',
    'print("\\nXGBoost vs LightGBM 汇总：")',
    'print(summary_df.to_string(index=False))',
])

# 保存
out_path = r'd:\cc_test\NeSI_XGBoost教程.docx'
doc.save(out_path)
print(f'Word 文档已生成：{out_path}')
